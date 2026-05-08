"""Generate ShipSmart Study Guide as a .docx file.

Run:  python scripts/build_study_guide.py
Output: docs/assets/ShipSmart-Study-Guide-v3.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def H1(doc, text):
    p = doc.add_heading(text, level=1)
    return p

def H2(doc, text):
    doc.add_heading(text, level=2)

def H3(doc, text):
    doc.add_heading(text, level=3)

def P(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def CODE(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(8)

def BULLET(doc, text):
    doc.add_paragraph(text, style="List Bullet")

def NUM(doc, text):
    doc.add_paragraph(text, style="List Number")

def PAGEBREAK(doc):
    doc.add_page_break()


doc = Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ── Title ──────────────────────────────────────────────────────────
title = doc.add_heading("ShipSmart — System Documentation & Interview Study Guide", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph(
    "A deep, codebase-grounded learning guide for Senior Backend / "
    "AI-Integrated Systems interviews\n"
    "Covers: Polyglot architecture · RAG pipeline · MCP/Tooling · Task-based LLM routing"
)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in sub.runs:
    r.italic = True

doc.add_paragraph()
toc = doc.add_paragraph()
toc.add_run("Contents\n").bold = True
toc.add_run(
    "1.  System Overview (Big Picture)\n"
    "2.  End-to-End Request Flow\n"
    "3.  Spring Boot Backend (Deep Dive)\n"
    "4.  FastAPI AI Service (Deep Dive)\n"
    "5.  RAG Concepts — Interview Focused\n"
    "6.  MCP / Tooling Concepts\n"
    "7.  LLM Design Decisions\n"
    "8.  Key Interview Talking Points\n"
    "9.  Common Pitfalls + What We Did Right\n"
    "10. 7-Day Study Plan\n"
    "11. Simplified Revision Notes\n"
)
PAGEBREAK(doc)


# ─────────────────────────────────────────────────────────────────────
# 1. SYSTEM OVERVIEW
# ─────────────────────────────────────────────────────────────────────
H1(doc, "1. System Overview (Big Picture)")

H2(doc, "1.1 What ShipSmart Is")
P(doc,
  "ShipSmart is a shipping comparison and management platform. Users "
  "search for shipping quotes, save options, get AI-powered shipping "
  "advice, and track bookings. The product question that drives the "
  "architecture is: 'How do we combine reliable transactional logic "
  "(quotes, bookings) with AI-assisted guidance (advice, "
  "recommendations) without letting one compromise the other?'")

H2(doc, "1.2 Polyglot 3-Service Architecture")
P(doc,
  "ShipSmart is a polyglot monorepo (Nx 22.3 + pnpm) containing three "
  "deployable services:")
BULLET(doc, "Frontend (apps/web) — React 19 + Vite + TypeScript + shadcn/ui + Tailwind. Static site on Render.")
BULLET(doc, "Java API (apps/api-java) — Spring Boot 4.0.5 (Java 25). Owns transactional data: quotes, saved options, booking redirects. Supabase PostgreSQL + JWT auth.")
BULLET(doc, "Python API (apps/api-python) — FastAPI 0.135.3 (Python 3.13). Owns AI features: RAG pipeline, LLM integration, tool orchestration, advisor endpoints, recommendations.")
P(doc,
  "Supporting infrastructure: Supabase (Postgres + Auth) and a small "
  "set of legacy edge functions kept feature-flagged as a fallback.")

H2(doc, "1.3 High-Level Architecture Diagram")
CODE(doc, r"""
                        ┌──────────────────────────┐
                        │   React 19 Frontend      │
                        │   (apps/web)             │
                        │   - Quote form           │
                        │   - Saved options        │
                        │   - Advisor page         │
                        └────────────┬─────────────┘
                                     │  HTTPS / JSON
                ┌────────────────────┴────────────────────┐
                │                                         │
       ┌────────▼─────────┐                     ┌─────────▼────────┐
       │  Java API        │                     │  Python API      │
       │  Spring Boot     │                     │  FastAPI         │
       │  apps/api-java   │                     │  apps/api-python │
       │                  │                     │                  │
       │  Owns:           │                     │  Owns:           │
       │  • Quotes        │                     │  • RAG pipeline  │
       │  • Saved options │                     │  • Tool registry │
       │  • Bookings      │                     │  • Advisors      │
       │  • JWT auth      │                     │  • LLM router    │
       └────────┬─────────┘                     └────┬─────────┬───┘
                │                                    │         │
       ┌────────▼─────────┐               ┌──────────▼──┐  ┌───▼─────────┐
       │   Supabase       │               │ Carrier APIs│  │  LLM APIs   │
       │   Postgres+Auth  │               │ UPS/FedEx/  │  │  OpenAI     │
       └──────────────────┘               │ DHL/USPS    │  │  Gemini     │
                                          │ (or Mock)   │  │  (Echo fb)  │
                                          └─────────────┘  └─────────────┘
""")
P(doc, "Important: the Java and Python APIs do not call each other. The "
       "frontend orchestrates between them. This is intentional — each "
       "service is independently deployable and the contract surface "
       "between them is the frontend's data shape, not an internal RPC.")

H2(doc, "1.4 Why Split Java and Python?")
P(doc, "The split is driven by ownership, not language preference:")
BULLET(doc, "Spring Boot is the system of record. It must be transactional, validated, auditable, and JWT-authenticated. Java/Spring is a mature, boring choice that nails this.")
BULLET(doc, "FastAPI hosts AI workloads where the ecosystem (LangChain-style libraries, OpenAI/Gemini SDKs, embeddings, vector libs) is Python-first. Mixing async I/O for LLM and HTTP calls is also natural in FastAPI.")
BULLET(doc, "Failure isolation: a flaky LLM provider or a poisoned RAG document cannot bring down the quoting/booking flow. If the Python API is down, the frontend simply hides the advisor card.")
BULLET(doc, "Independent deploy and scale: the AI service is more volatile (model swaps, prompt tweaks). Keeping it on its own deploy cycle protects the transactional service.")

PAGEBREAK(doc)


# ─────────────────────────────────────────────────────────────────────
# 2. END-TO-END FLOW
# ─────────────────────────────────────────────────────────────────────
H1(doc, "2. End-to-End Request Flow")

H2(doc, "2.1 Scenario: 'Best shipping option for a fragile item'")
P(doc, "Concrete walk-through of what happens, layer by layer, when the "
       "user types this question into the Advisor page.")

CODE(doc, r"""
┌──────────────┐    POST /api/v1/advisor/shipping
│  Frontend    │ ───────────────────────────────►
│ AdvisorPage  │      { query, context }
└──────────────┘
                                  │
                                  ▼
                ┌───────────────────────────────────┐
                │ FastAPI route: advisor.shipping   │
                │ apps/api-python/app/api/routes/   │
                │ advisor.py:32                     │
                └─────────────────┬─────────────────┘
                                  │ pull rag + tool_registry +
                                  │ llm_router from app.state
                                  ▼
                ┌───────────────────────────────────┐
                │ get_shipping_advice()             │
                │ services/shipping_advisor_service │
                └─────────────────┬─────────────────┘
                                  │
        ┌─────────────────────────┼─────────────────────────┐
        ▼                         ▼                         ▼
 ┌──────────────┐        ┌────────────────┐        ┌───────────────┐
 │ RAG retrieve │        │ Tool execution │        │ LLM reasoning │
 │ embed query →│        │ (only when     │        │ via router    │
 │ vector search│        │ context flags  │        │ for_task(     │
 │ top_k chunks │        │ trigger them)  │        │  REASONING)   │
 └──────┬───────┘        └────────┬───────┘        └───────┬───────┘
        │                         │                        │
        └─────────────────────────┼────────────────────────┘
                                  ▼
                ┌───────────────────────────────────┐
                │ build_advisor_prompt( query,      │
                │   context = retrieved + tools )   │
                └─────────────────┬─────────────────┘
                                  ▼
                ┌───────────────────────────────────┐
                │ OpenAI Chat Completions           │
                │ (reasoning task → openai)         │
                └─────────────────┬─────────────────┘
                                  ▼
                ┌───────────────────────────────────┐
                │ ShippingAdvisorResponse JSON      │
                │ { answer, reasoning_summary,      │
                │   tools_used, sources,            │
                │   context_used }                  │
                └─────────────────┬─────────────────┘
                                  ▼
                       ┌────────────────────┐
                       │  Frontend renders  │
                       │  in AdvisorCard    │
                       └────────────────────┘
""")

H2(doc, "2.2 Step-by-Step Narration")
NUM(doc, "User types the query and hits submit on AdvisorPage. The frontend POSTs to /api/v1/advisor/shipping on the Python API. The Java API is not involved — this is a pure AI flow.")
NUM(doc, "FastAPI route handler (advisor.py) pulls three things off app.state: rag (embedding_provider, vector_store), tool_registry, and llm_router. These were built once at startup by the lifespan hook in main.py.")
NUM(doc, "The route asks the router for the reasoning client: llm_router.for_task(TASK_REASONING). With the recommended config that returns the OpenAI client.")
NUM(doc, "get_shipping_advice() runs the RAG retrieval: embed the query, cosine-search the in-memory vector store, return the top-K chunks with their source paths and scores.")
NUM(doc, "If the request context contains hints (e.g. an address dict), the service runs ValidateAddressTool / GetQuotePreviewTool from the registry. Tool outputs are appended to the prompt context.")
NUM(doc, "build_advisor_prompt() composes a single prompt with the original query, the retrieved chunks, and the tool outputs. The prompt is grounded — the system message says 'answer ONLY from the provided context'.")
NUM(doc, "The OpenAI client sends the messages to gpt-4o-mini with configured timeout, max_tokens, and temperature. On any error, AppError(502) is raised by the client and the global error handler converts it to a clean JSON response.")
NUM(doc, "The service builds a structured response: answer text, reasoning summary, tools_used list, sources list (file path + chunk index + score), and context_used flag.")
NUM(doc, "The frontend renders the structured fields. Sources become citation chips so the user can see exactly which document the answer came from.")

H2(doc, "2.3 Contrast: RAG-Only Query Flow")
P(doc, "POST /api/v1/rag/query is simpler — no tool layer. The route "
       "calls llm_router.for_task(TASK_SYNTHESIS) which returns the "
       "Gemini client. Retrieval → prompt → Gemini → return answer + "
       "sources. Same retrieval code, different LLM, different prompt "
       "(build_rag_prompt vs build_advisor_prompt).")

PAGEBREAK(doc)


# ─────────────────────────────────────────────────────────────────────
# 3. SPRING BOOT BACKEND
# ─────────────────────────────────────────────────────────────────────
H1(doc, "3. Spring Boot Backend (Deep Dive)")

H2(doc, "3.1 Role")
P(doc, "Spring Boot is the system of record for everything money-shaped "
       "or audit-shaped: the user's quote requests, their saved options, "
       "and their booking redirects. Anything that has to be correct "
       "even if the AI service is on fire lives here.")

H2(doc, "3.2 Key Endpoints")
BULLET(doc, "POST /api/v1/quotes — submit a 3-step quote form, get carrier comparison.")
BULLET(doc, "CRUD /api/v1/saved-options — JWT-protected, per-user saved quotes.")
BULLET(doc, "POST /api/v1/bookings/redirect — record that the user clicked through to a carrier site to book.")

H2(doc, "3.3 Why Java for the Transactional Layer?")
BULLET(doc, "Mature transactional story: Spring Data JPA, declarative @Transactional, validated DTOs, global exception handlers.")
BULLET(doc, "Strong type system catches schema/contract drift early — important for the data the business cares about.")
BULLET(doc, "JWT auth with Spring Security is well-trodden ground; tokens come from Supabase Auth and are validated server-side.")
BULLET(doc, "It is the boring choice. The right amount of boring for the part of the system that must never lose data.")

H2(doc, "3.4 How Java Interacts With Python")
P(doc, "It does not. By design. The frontend calls each service "
       "directly and merges results in the UI. The two backends share "
       "nothing except the user identity (Supabase JWT). This means:")
BULLET(doc, "No circular dependency between services.")
BULLET(doc, "Independent deploys and rollbacks.")
BULLET(doc, "An LLM outage cannot block a quote being saved.")
BULLET(doc, "Each backend has exactly one persistence story.")

PAGEBREAK(doc)


# ─────────────────────────────────────────────────────────────────────
# 4. FASTAPI AI SERVICE
# ─────────────────────────────────────────────────────────────────────
H1(doc, "4. FastAPI AI Service (Deep Dive)")

H2(doc, "4.1 Core Architecture")
P(doc, "The Python service is built around a single principle: build "
       "everything heavy once at startup, hang it off app.state, and "
       "let route handlers borrow it. This makes request handling thin "
       "and testable.")

H3(doc, "Lifecycle (apps/api-python/app/main.py)")
CODE(doc, r"""
@asynccontextmanager
async def lifespan(app: FastAPI):
    # 1. Shared httpx client
    app.state.http_client = httpx.AsyncClient(...)

    # 2. RAG components
    embedding_provider = create_embedding_provider()
    vector_store       = create_vector_store()

    # 3. LLM router (one client per task)
    llm_router = create_llm_router()
    app.state.llm_router = llm_router

    # 4. Back-compat shim — synthesis client also under app.state.rag
    app.state.rag = {
        "embedding_provider": embedding_provider,
        "vector_store":       vector_store,
        "llm_client":         llm_router.for_task(TASK_SYNTHESIS),
    }

    # 5. Tool registry bound to a shipping provider
    shipping_provider = create_shipping_provider()
    tool_registry = ToolRegistry()
    tool_registry.register(ValidateAddressTool(shipping_provider))
    tool_registry.register(GetQuotePreviewTool(shipping_provider))
    app.state.tool_registry = tool_registry

    yield

    await app.state.http_client.aclose()
""")

H3(doc, "Config System")
P(doc, "app/core/config.py uses pydantic-settings. Every env var has a "
       "typed field with a default, so the app boots even with an empty "
       ".env. Properties like cors_origins_list and is_production keep "
       "the route handlers free of string-parsing.")

H2(doc, "4.2 RAG Pipeline")

H3(doc, "Folder layout of knowledge base")
CODE(doc, r"""
apps/api-python/data/documents/
├── carriers/   (UPS, FedEx, DHL, USPS overviews)
├── guides/     (FAQ, ground vs express, packaging, address quality)
├── policies/   (returns/claims, international, comparisons)
└── scenarios/  (recommendation tradeoffs, delays/exceptions)
""")

H3(doc, "Ingestion (app/rag/ingestion.py)")
NUM(doc, "load_documents(directory) walks **/*.txt and **/*.md recursively. The relative path becomes the source name (e.g. carriers/ups-overview.md), so retrieval results are self-locating.")
NUM(doc, "chunk_text(content, chunk_size=500, chunk_overlap=50) splits each document into overlapping windows with their source + index.")
NUM(doc, "embedding_provider.embed(texts) turns chunk text into vectors. With EMBEDDING_PROVIDER='openai' it uses text-embedding-3-small; otherwise LocalHashEmbedding (deterministic, no semantics — fine for dev).")
NUM(doc, "vector_store.add(stored_chunks) persists into InMemoryVectorStore. zip(strict=True) catches any chunk/embedding count mismatch loudly.")

H3(doc, "Retrieval")
P(doc, "rag_query() embeds the user query, calls vector_store.search() "
       "(cosine similarity, top_k from settings.rag_top_k), then passes "
       "the chunks to build_rag_prompt() and the LLM. The response "
       "carries the sources back to the caller so the UI can cite them.")

H2(doc, "4.3 MCP / Tooling Layer")
P(doc, "Tools are deterministic, schema-validated functions that the "
       "LLM (or our own service code) can call. They exist so that "
       "anything with a right answer — address validation, quote "
       "lookup — is computed by code, not hallucinated by the model.")

CODE(doc, r"""
class Tool(ABC):
    name: str
    description: str
    parameters: list[ParameterSchema]
    async def execute(self, **kwargs) -> ToolResult: ...
    def schema(self) -> dict: ...
    def validate_input(self, kwargs) -> None: ...

class ToolRegistry:
    register(tool)
    get(name) -> Tool
    list_tools() -> list[Tool]
    list_schemas() -> list[dict]   # for /orchestration/tools
""")

P(doc, "ValidateAddressTool and GetQuotePreviewTool both delegate to "
       "the configured ShippingProvider, so tools work with mock or "
       "real carriers without code changes.")

H2(doc, "4.4 Provider Integration")
P(doc, "ShippingProvider is an ABC with validate_address() and "
       "get_quote_preview(). create_shipping_provider() reads "
       "SHIPPING_PROVIDER from config, dynamically imports the matching "
       "module, validates required credentials, and falls back to "
       "MockShippingProvider on any failure. The app never crashes "
       "because of carrier env-var problems.")

CODE(doc, r"""
SHIPPING_PROVIDER=ups   →  UPSProvider   (needs UPS_CLIENT_ID + SECRET)
SHIPPING_PROVIDER=fedex →  FedExProvider (needs FEDEX_CLIENT_ID + SECRET)
SHIPPING_PROVIDER=dhl   →  DHLProvider   (needs DHL_API_KEY + SECRET)
SHIPPING_PROVIDER=usps  →  USPSProvider  (needs USPS_CLIENT_ID + SECRET)
SHIPPING_PROVIDER=mock  →  MockShippingProvider
unset / unknown / bad creds → MockShippingProvider (with warning)
""")

H2(doc, "4.5 LLM Integration")
P(doc, "LLMClient is an ABC with one method: complete(messages) -> "
       "str. Concrete implementations:")
BULLET(doc, "OpenAIClient — AsyncOpenAI SDK, configurable model/timeout/temperature/max_tokens. Maps SDK errors to AppError(502).")
BULLET(doc, "GeminiClient — REST via httpx. Includes a _messages_to_gemini_contents helper that merges the OpenAI-style system message into the first user turn (Gemini has no system role).")
BULLET(doc, "LlamaClient — talks to a local Ollama server using the OpenAI-compatible /v1 endpoint. Useful for offline/dev.")
BULLET(doc, "EchoClient — placeholder. Returns retrieved context as text. Used as the universal safety net.")

H2(doc, "4.6 Task-Based LLM Routing (the Phase 14 work)")
P(doc, "The router (app/llm/router.py) maps a logical task to a "
       "concrete LLMClient. Three task names exist:")
BULLET(doc, "TASK_REASONING — used by /advisor/shipping and /advisor/tracking. Reasoning over RAG context + tool results, where instruction-following matters more than cost.")
BULLET(doc, "TASK_SYNTHESIS — used by /rag/query and /advisor/recommendation. Grounded text generation from already-structured input. Cheaper, faster models are fine here.")
BULLET(doc, "TASK_FALLBACK — the safety net. Always points at something that always works (Echo by default).")

H3(doc, "Resolution order per task")
NUM(doc, "LLM_PROVIDER_<TASK> if set (e.g. LLM_PROVIDER_REASONING=openai)")
NUM(doc, "Legacy LLM_PROVIDER (so older deployments keep working)")
NUM(doc, "LLM_PROVIDER_FALLBACK (the configured fallback provider)")
NUM(doc, "EchoClient (always available)")

CODE(doc, r"""
# Recommended production config
LLM_PROVIDER_REASONING=openai   OPENAI_API_KEY=sk-...
LLM_PROVIDER_SYNTHESIS=gemini   GEMINI_API_KEY=...
LLM_PROVIDER_FALLBACK=echo

# Resulting router on startup:
# LLM router initialized: {
#   'reasoning': 'openai',
#   'synthesis': 'gemini',
#   'fallback': 'echo'
# }
""")

P(doc, "Crucially, each task degrades independently. A missing OpenAI "
       "key only collapses the reasoning task to echo — synthesis with "
       "Gemini keeps working. This is verified by "
       "test_missing_key_falls_back_to_fallback_provider.")

PAGEBREAK(doc)


# ─────────────────────────────────────────────────────────────────────
# 5. RAG CONCEPTS
# ─────────────────────────────────────────────────────────────────────
H1(doc, "5. RAG Concepts (Interview-Focused)")

H2(doc, "5.1 What is RAG?")
P(doc, "Retrieval-Augmented Generation. Instead of asking the LLM to "
       "answer from its frozen training data, you (1) retrieve relevant "
       "passages from your own knowledge base, (2) put them in the "
       "prompt, (3) instruct the model to answer ONLY from those "
       "passages. The model becomes a reading-comprehension engine over "
       "your data, not a memory engine.")

H2(doc, "5.2 Why Not Just Use the LLM Alone?")
BULLET(doc, "Hallucination: LLMs invent confident answers when they don't know.")
BULLET(doc, "Stale knowledge: training cutoffs are months or years old.")
BULLET(doc, "No private data: an LLM has never seen your carrier policies.")
BULLET(doc, "No citations: free-form generation can't tell users where the answer came from.")
P(doc, "RAG fixes all four with one design move: ground every answer "
       "in retrieved text and surface that text as a citation.")

H2(doc, "5.3 Chunking")
P(doc, "Why not embed whole documents? Two reasons. (1) The model has "
       "a context limit. (2) Long documents dilute the embedding — a "
       "single vector can't represent twenty different topics. We chunk "
       "into ~500-character windows with 50-character overlap. The "
       "overlap means a sentence that straddles a boundary still "
       "appears in at least one chunk in full.")
P(doc, "In this codebase: app/rag/chunking.py + RAG_CHUNK_SIZE / "
       "RAG_CHUNK_OVERLAP env vars.")

H2(doc, "5.4 Embeddings")
P(doc, "An embedding turns text into a fixed-length vector such that "
       "semantically similar text lands at nearby points. 'How long "
       "does ground shipping take?' and 'UPS Ground delivery time' end "
       "up close together even though they share almost no words.")
BULLET(doc, "Production: OpenAIEmbedding with text-embedding-3-small (256-dim).")
BULLET(doc, "Dev/CI: LocalHashEmbedding — deterministic hash-based vectors. No real semantics; lets the pipeline run without external calls.")

H2(doc, "5.5 Vector Search")
P(doc, "Cosine similarity between the query embedding and every stored "
       "chunk embedding. Top-K (default 3) wins. In production you'd "
       "use a real ANN index (pgvector, Qdrant, FAISS); we use "
       "InMemoryVectorStore with a brute-force scan because the corpus "
       "is small (~150 chunks) and the goal is to demonstrate the "
       "shape of the system, not the index.")

H2(doc, "5.6 The Retrieve-Then-Read Flow")
CODE(doc, r"""
query
  │
  ▼
embed(query) ─────►  vector
                       │
                       ▼
              vector_store.search(vector, top_k=3)
                       │
                       ▼
            top_k chunks (text + source + score)
                       │
                       ▼
       build_rag_prompt(query, [chunk1, chunk2, chunk3])
                       │
                       ▼
              LLM.complete(messages)
                       │
                       ▼
                grounded answer + sources
""")

PAGEBREAK(doc)


# ─────────────────────────────────────────────────────────────────────
# 6. MCP / TOOLING
# ─────────────────────────────────────────────────────────────────────
H1(doc, "6. MCP / Tooling Concepts")

H2(doc, "6.1 Why LLMs Need Tools")
P(doc, "An LLM is a text predictor. Ask it 'is 90210 a valid ZIP code "
       "in California?' and you'll often get a confident yes — even "
       "for ZIPs that don't exist. Ask it 'what does UPS Ground cost "
       "from Boston to Seattle for a 5lb package?' and it will invent "
       "a number. Tools fix this by routing those questions to code.")

H2(doc, "6.2 Free Text vs Structured Execution")
CODE(doc, r"""
                  ┌─────────────────────────────────────┐
                  │  LLM (text in, text out)            │
                  └─────┬───────────────────────────────┘
                        │
   ┌────────────────────┼────────────────────┐
   ▼                    ▼                    ▼
free text         structured tool      structured tool
generation        call (validated)     call (validated)
   │                    │                    │
   ▼                    ▼                    ▼
'1-5 days'         ValidateAddress      GetQuotePreview
'usually fine'     → real provider      → real provider
                   → real result        → real result
""")

P(doc, "The free-text path is fine for tone and reasoning. The "
       "structured path is where ground truth lives. A good system "
       "uses both, with clear handoffs.")

H2(doc, "6.3 How Our Tool Layer Works")
NUM(doc, "Each tool subclasses Tool with name, description, parameters, and async execute().")
NUM(doc, "ToolRegistry holds them. /api/v1/orchestration/tools returns the JSON schemas — that is the MCP-style discovery surface.")
NUM(doc, "ValidateAddressTool and GetQuotePreviewTool delegate to the active ShippingProvider, which is mock or real depending on env vars.")
NUM(doc, "Inputs are validated before execute() runs. Errors raise AppError so the global handler maps them to clean HTTP responses.")
NUM(doc, "Tool outputs are appended to the LLM prompt as evidence so the model's answer is grounded in real data, not invented.")

H2(doc, "6.4 How This Avoids Hallucination")
BULLET(doc, "Address questions go to ValidateAddressTool, not to the LLM.")
BULLET(doc, "Pricing/ETA questions go to GetQuotePreviewTool, not to the LLM.")
BULLET(doc, "Open-ended advice goes to the LLM, but with retrieved policy text in the prompt and an instruction to answer only from that text.")
BULLET(doc, "Sources are returned to the UI so users can verify any answer.")

PAGEBREAK(doc)


# ─────────────────────────────────────────────────────────────────────
# 7. LLM DESIGN DECISIONS
# ─────────────────────────────────────────────────────────────────────
H1(doc, "7. LLM Design Decisions")

H2(doc, "7.1 Why Multiple LLMs?")
P(doc, "Different tasks have different cost/quality curves. Reasoning "
       "over tool results and policy snippets benefits from a stronger "
       "instruction-following model. Synthesizing a short summary from "
       "already-structured data is a much simpler job and a cheaper "
       "model handles it indistinguishably. Mixing providers also "
       "removes single-vendor risk.")

H2(doc, "7.2 Why Task-Based Routing (Not a Single Model)?")
BULLET(doc, "Cost: Gemini Flash is roughly an order of magnitude cheaper than OpenAI per token. Sending light synthesis to it cuts the LLM bill without hurting users.")
BULLET(doc, "Failure isolation: a Gemini outage breaks RAG q&a but the advisor reasoning path stays up.")
BULLET(doc, "Operability: ops can swap a task's provider via env var only. Zero code change, zero new endpoint.")
BULLET(doc, "Honest interview answer: it's also a clean abstraction to talk about — task → provider, with explicit fallback. Easy to whiteboard.")

H2(doc, "7.3 Why OpenAI for Reasoning?")
BULLET(doc, "Strong instruction-following on grounded prompts.")
BULLET(doc, "Reliable JSON / structured output behaviour.")
BULLET(doc, "gpt-4o-mini is the cheapest model in the family that handles our prompts reliably; we don't need 4o for this workload.")

H2(doc, "7.4 Why Gemini for Synthesis?")
BULLET(doc, "Roughly 10× cheaper per token than OpenAI's small models.")
BULLET(doc, "Quality difference on grounded summarisation is small to nil.")
BULLET(doc, "Latency on Flash is noticeably lower for short outputs.")

H2(doc, "7.5 Why a Fallback at All?")
P(doc, "EchoClient never makes a network call and never fails. Wiring "
       "it as the safety net guarantees the API stays up — and "
       "interviews can demo end-to-end — when keys are missing or both "
       "providers are down.")

H2(doc, "7.6 Tradeoffs")
BULLET(doc, "Two providers means two SDK surfaces, two error modes, two prompt quirks (Gemini has no system role). The router contains this complexity in one file.")
BULLET(doc, "Per-task tuning (timeout/temperature/max_tokens) is shared across tasks today — that's a known limitation, easy to extend.")
BULLET(doc, "Mixed providers can produce slightly different tone between reasoning and synthesis tasks. Acceptable here because each task talks to a different UI surface.")

PAGEBREAK(doc)


# ─────────────────────────────────────────────────────────────────────
# 8. INTERVIEW TALKING POINTS
# ─────────────────────────────────────────────────────────────────────
H1(doc, "8. Key Interview Talking Points")

H2(doc, "8.1 'Explain your system architecture.'")
P(doc, "ShipSmart is a polyglot monorepo with three deployable services. "
       "A React frontend orchestrates between two backends: a Spring "
       "Boot service that owns transactional data (quotes, saved "
       "options, bookings) backed by Supabase Postgres with JWT auth, "
       "and a FastAPI service that owns AI features — RAG, tools, "
       "advisors, and a task-based LLM router. The two backends never "
       "call each other; they share only the user's JWT identity. "
       "That gives independent deployability and means an LLM outage "
       "cannot block a quote being saved.")

H2(doc, "8.2 'How does your RAG pipeline work?'")
P(doc, "On startup we walk data/documents recursively, chunk each file "
       "into ~500-char overlapping windows, embed them — OpenAI "
       "text-embedding-3-small in production, a deterministic hash "
       "embedding in dev — and store them in an in-memory vector "
       "store. At query time we embed the query, take cosine top-3, "
       "and call build_rag_prompt() which puts the chunks in a system-"
       "grounded prompt that tells the LLM to answer only from the "
       "context. The endpoint returns the answer plus a list of "
       "sources (file path + chunk index + similarity score) so the "
       "UI can cite them.")

H2(doc, "8.3 'How do you avoid hallucination?'")
NUM(doc, "Grounding: the system prompt explicitly says 'answer ONLY from the provided context'.")
NUM(doc, "Retrieval over generation: we put real document text into the prompt instead of trusting the model's parametric memory.")
NUM(doc, "Tools for ground-truth questions: addresses and quotes go through ValidateAddressTool / GetQuotePreviewTool, not the LLM.")
NUM(doc, "Citations: we return sources to the UI so the user can verify any answer.")
NUM(doc, "Echo fallback: if the LLM is unavailable we return the retrieved context verbatim rather than nothing.")

H2(doc, "8.4 'Why separate Java and Python services?'")
P(doc, "Ownership and failure isolation, not language preference. "
       "Spring is the boring, mature choice for transactional data — "
       "JPA, @Transactional, validated DTOs, JWT auth. FastAPI is "
       "where the AI ecosystem lives — OpenAI/Gemini SDKs, async I/O, "
       "embedding libraries. Splitting them means the AI service can "
       "deploy faster, fail independently, and never block the "
       "quoting flow.")

H2(doc, "8.5 'How would you scale this system?'")
BULLET(doc, "Replace InMemoryVectorStore with pgvector or Qdrant — same VectorStore ABC, swap one factory.")
BULLET(doc, "Add a real ANN index — brute force is fine at 150 chunks but breaks down past ~10k.")
BULLET(doc, "Cache embeddings of frequent queries (Redis) and cache LLM completions for identical prompts.")
BULLET(doc, "Move ingestion offline — currently it's a route; in production it would be a background job.")
BULLET(doc, "Add per-task timeouts and circuit breakers around each LLM provider.")
BULLET(doc, "Horizontal scale FastAPI behind a load balancer; the service is stateless apart from app.state caches that are easy to externalise.")

H2(doc, "8.6 'How would you improve this system?'")
BULLET(doc, "Per-task LLM tuning (LLM_TEMPERATURE_SYNTHESIS, LLM_MODEL_REASONING).")
BULLET(doc, "Persistent vector store with metadata filtering.")
BULLET(doc, "Function calling / proper MCP server interface so external clients can discover and call our tools.")
BULLET(doc, "Multi-turn advisor conversations with history.")
BULLET(doc, "Rate limiting and auth on the Python API.")
BULLET(doc, "Eval harness for RAG quality (recall@k, faithfulness scoring).")

PAGEBREAK(doc)


# ─────────────────────────────────────────────────────────────────────
# 9. PITFALLS
# ─────────────────────────────────────────────────────────────────────
H1(doc, "9. Common Pitfalls + What We Did Right")

H2(doc, "9.1 Pitfalls People Hit Building RAG")
BULLET(doc, "Embedding whole documents — breaks retrieval quality. Fix: chunk with overlap.")
BULLET(doc, "Letting the LLM answer from memory when context is empty — leads to hallucination. Fix: prompt explicitly says 'no context → say so'.")
BULLET(doc, "No citations — users can't verify the answer. Fix: pass sources back from retrieval to API response.")
BULLET(doc, "Single global LLM client — one outage takes everything down. Fix: per-task router with independent fallback.")
BULLET(doc, "Hardcoded provider — no way to swap models without code change. Fix: env-driven factory + ABC.")
BULLET(doc, "Tools that re-implement business logic — drift between code and what the LLM 'knows'. Fix: tools delegate to the same provider abstraction the rest of the system uses.")
BULLET(doc, "Crashing the app on missing env vars — fragile. Fix: every factory degrades to a safe default with a warning.")

H2(doc, "9.2 What This Codebase Does Right")
BULLET(doc, "Clean ABCs (LLMClient, EmbeddingProvider, VectorStore, ShippingProvider, Tool) — implementations are interchangeable.")
BULLET(doc, "Single startup, hang-on-app.state pattern — request handlers are thin and testable.")
BULLET(doc, "Task-based routing built on a config-driven resolver — no heuristics, no hidden behaviour.")
BULLET(doc, "Per-task fallback so partial outages stay partial.")
BULLET(doc, "Recursive document loader using the relative path as the source name — citations are self-locating.")
BULLET(doc, "171 Python tests including the router/factory/missing-key cases. The code is pinned by tests, not vibes.")

PAGEBREAK(doc)


# ─────────────────────────────────────────────────────────────────────
# 10. STUDY PLAN
# ─────────────────────────────────────────────────────────────────────
H1(doc, "10. 7-Day Study Plan")

P(doc, "Each day pairs reading with a hands-on action and an "
       "out-loud explanation. Speaking the answer is what locks it "
       "in for the interview.")

H2(doc, "Day 1 — System Overview & Architecture")
BULLET(doc, "Read: docs/current-system-state.md, this guide §1 and §2.")
BULLET(doc, "Run: start both backends locally; hit /health on each.")
BULLET(doc, "Understand: the three services and the rule that Java and Python don't talk to each other.")
BULLET(doc, "Out loud: explain the architecture in 90 seconds without looking at the diagram.")

H2(doc, "Day 2 — Spring Boot Layer")
BULLET(doc, "Read: apps/api-java controllers, services, and security config.")
BULLET(doc, "Run: POST a quote, save an option, hit the booking redirect endpoint.")
BULLET(doc, "Understand: why transactional data lives here, not in Python.")
BULLET(doc, "Out loud: explain why splitting Java and Python is a feature not a cost.")

H2(doc, "Day 3 — FastAPI Skeleton & Lifecycle")
BULLET(doc, "Read: apps/api-python/app/main.py, app/core/config.py.")
BULLET(doc, "Run: uv run uvicorn app.main:app --reload, watch the startup logs (LLM router init line).")
BULLET(doc, "Understand: the lifespan() pattern and how app.state holds prebuilt clients.")
BULLET(doc, "Out loud: explain why building things once at startup is better than per-request construction.")

H2(doc, "Day 4 — RAG Pipeline")
BULLET(doc, "Read: app/rag/{ingestion,chunking,embeddings,vector_store}.py + app/services/rag_service.py.")
BULLET(doc, "Run: POST /api/v1/rag/ingest, then POST /api/v1/rag/query with a real question.")
BULLET(doc, "Understand: ingest → chunk → embed → store → retrieve → prompt → generate.")
BULLET(doc, "Out loud: explain RAG to someone who has never heard the term, in two minutes.")

H2(doc, "Day 5 — Tools / MCP Layer")
BULLET(doc, "Read: app/tools/{base,registry,address_tools,quote_tools}.py + app/api/routes/orchestration.py.")
BULLET(doc, "Run: GET /api/v1/orchestration/tools to see the schema list. POST /api/v1/orchestration/run with an explicit tool name.")
BULLET(doc, "Understand: why deterministic tools sit alongside the LLM and how they avoid hallucination on factual questions.")
BULLET(doc, "Out loud: explain the 'free text vs structured execution' split using ValidateAddressTool as the example.")

H2(doc, "Day 6 — LLM Integration & Task-Based Routing")
BULLET(doc, "Read: app/llm/client.py, app/llm/router.py, app/llm/prompts.py.")
BULLET(doc, "Run: try LLM_PROVIDER_REASONING=openai LLM_PROVIDER_SYNTHESIS=gemini and watch the router init log. Then break a key on purpose and confirm only one task degrades.")
BULLET(doc, "Understand: per-task resolution order, the role of EchoClient, and why the abstraction sits at the ABC layer.")
BULLET(doc, "Out loud: explain why a single LLM client would have been worse and walk through the resolution order from memory.")

H2(doc, "Day 7 — Provider Layer + Frontend Integration + Mock Interview")
BULLET(doc, "Read: app/providers/{__init__,shipping_provider,mock_provider,ups_provider}.py + apps/web src AdvisorPage and RecommendationCard.")
BULLET(doc, "Run: SHIPPING_PROVIDER=ups with empty creds → confirm fallback to mock; then SHIPPING_PROVIDER=mock and trace a full advisor request from the UI.")
BULLET(doc, "Understand: how the same Tool / Service / LLM stack works against mock and real carriers.")
BULLET(doc, "Out loud: do all six §8 talking points without notes. Time them.")

PAGEBREAK(doc)


# ─────────────────────────────────────────────────────────────────────
# 11. SIMPLIFIED REVISION
# ─────────────────────────────────────────────────────────────────────
H1(doc, "11. Simplified Revision Notes")

H2(doc, "11.1 One-Page Mental Model")
CODE(doc, r"""
ShipSmart = 3 services:
  Frontend (React)  →  React 19, talks to BOTH backends
  Java API          →  transactional (quotes, saved, bookings)
  Python API        →  AI (RAG, tools, advisors, LLM router)

Java and Python NEVER call each other. Frontend orchestrates.

Python API has 5 layers:
  config       → pydantic-settings, env-driven
  providers    → ShippingProvider ABC, mock + 4 carriers
  rag          → ingest → chunk → embed → store → retrieve
  tools        → ToolRegistry, deterministic, schema-validated
  llm          → LLMClient ABC + LLMRouter (task → client)

Task-based routing:
  reasoning  → OpenAI (advisor endpoints)
  synthesis  → Gemini (rag/query, recommendation)
  fallback   → Echo (always works)

Resolution: LLM_PROVIDER_<TASK> → LLM_PROVIDER → fallback → Echo
""")

H2(doc, "11.2 30-Second Pitch")
P(doc, "ShipSmart is a shipping comparison platform. It splits "
       "transactional logic (Spring Boot, Postgres, JWT) from AI "
       "features (FastAPI with a RAG pipeline, a tool registry, and a "
       "task-based LLM router that points reasoning at OpenAI and "
       "synthesis at Gemini, with an Echo fallback so the system never "
       "fails on missing keys). The two backends never call each "
       "other — the frontend orchestrates — which gives independent "
       "deployability and failure isolation.")

H2(doc, "11.3 Quick-Reference Table")
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Concept"
hdr[1].text = "Where in code"
hdr[2].text = "One-line summary"

rows = [
    ("Lifecycle wiring", "app/main.py", "Build everything once at startup; hang on app.state."),
    ("Config", "app/core/config.py", "pydantic-settings, every field has a default."),
    ("RAG ingestion", "app/rag/ingestion.py", "Recursive walk → chunk → embed → store."),
    ("RAG query", "app/services/rag_service.py", "Embed query → top-K → grounded prompt → LLM."),
    ("Tool registry", "app/tools/registry.py", "Discoverable, schema-validated, deterministic."),
    ("Provider factory", "app/providers/__init__.py", "Env-driven, falls back to mock on bad creds."),
    ("LLM client ABC", "app/llm/client.py", "complete(messages) → str. Four implementations."),
    ("LLM router", "app/llm/router.py", "Task → client, resolved at startup, per-task fallback."),
    ("Reasoning task", "advisor.py routes", "OpenAI by default. Used by /advisor/shipping & /tracking."),
    ("Synthesis task", "rag.py + advisor.py", "Gemini by default. Used by /rag/query & /recommendation."),
]
for c, w, s in rows:
    cells = table.add_row().cells
    cells[0].text = c
    cells[1].text = w
    cells[2].text = s

doc.add_paragraph()
H2(doc, "11.4 Final Sanity Checks Before the Interview")
BULLET(doc, "Can I draw the architecture in under 60 seconds without looking?")
BULLET(doc, "Can I explain RAG without using the word 'RAG'?")
BULLET(doc, "Can I name the four LLMClient implementations and one reason each exists?")
BULLET(doc, "Can I walk the resolution order for the reasoning task from memory?")
BULLET(doc, "Can I name three things I'd improve and why?")
BULLET(doc, "Have I actually run /api/v1/rag/query and /api/v1/advisor/shipping in the last 24 hours?")

PAGEBREAK(doc)


# ─────────────────────────────────────────────────────────────────────
# 12. CODE TRACING MAPS
# ─────────────────────────────────────────────────────────────────────
H1(doc, "12. Code Tracing Maps")
P(doc, "Each map below traces a real request through every file it "
       "touches, in order. Open them side-by-side with the codebase "
       "and you can read along line by line. These are the maps to "
       "rehearse before a whiteboard interview.")

H2(doc, "12.1 Quote Flow (Java API)")
CODE(doc, r"""
Frontend (apps/web)
  └─ POST /api/v1/quotes  { origin, destination, package, ... }
        │
        ▼
apps/api-java
  ├─ controller/QuoteController.getQuote()
  │     - @Valid QuoteRequest dto         (Bean Validation runs here)
  │     - @PreAuthorize / SecurityFilter  (JWT extracted upstream)
  │
  ├─ service/QuoteService.computeQuotes()
  │     - business rules, normalisation
  │     - calls carrier comparison logic
  │     - @Transactional boundary opens
  │
  ├─ repository/QuoteRepository  (Spring Data JPA)
  │     - persists the quote request for audit
  │
  └─ returns QuoteResponse
        - list of carrier options
        - validation errors → 400 via @ControllerAdvice
        - auth failures     → 401/403 via SecurityFilterChain
        │
        ▼
Frontend renders the comparison table.
""")

H2(doc, "12.2 Advisor Flow (Python API)")
CODE(doc, r"""
Frontend AdvisorPage
  └─ POST /api/v1/advisor/shipping  { query, context }
        │
        ▼
apps/api-python/app/api/routes/advisor.py:32  shipping_advisor()
  - rag           = request.app.state.rag
  - tool_registry = request.app.state.tool_registry
  - llm_router    = request.app.state.llm_router
  - reasoning_client = llm_router.for_task(TASK_REASONING)   # → OpenAI
        │
        ▼
app/services/shipping_advisor_service.py  get_shipping_advice()
  ├─ embedding_provider.embed([query])
  ├─ vector_store.search(qvec, top_k=settings.rag_top_k)     # cosine
  ├─ if context has address dict:
  │     tool_registry.get("validate_address").execute(...)
  │     tool_registry.get("get_quote_preview").execute(...)
  ├─ build_advisor_prompt(query, retrieved_chunks, tool_results)
  └─ reasoning_client.complete(messages)                     # OpenAI
        │
        ▼
ShippingAdvisorResponse {
  answer, reasoning_summary, tools_used,
  sources [{source, chunk_index, score}],
  context_used
}
""")

H2(doc, "12.3 Recommendation Flow (Python API)")
CODE(doc, r"""
Frontend RecommendationCard (after a quote is fetched)
  └─ POST /api/v1/advisor/recommendation  { services, context }
        │
        ▼
app/api/routes/advisor.py:106  get_recommendation()
  - llm_router = request.app.state.llm_router
  - llm_client = llm_router.for_task(TASK_SYNTHESIS)         # → Gemini
        │
        ▼
app/services/recommendation_service.py  generate_recommendations()
  ├─ score each service deterministically
  │     (price, speed, reliability — pure code, no LLM)
  ├─ classify primary vs alternatives by score
  ├─ if llm_client present:
  │     summary = await _generate_summary(primary, alts, llm_client)
  │     (this is the ONLY place the LLM touches recommendations)
  └─ return RecommendationResult
        │
        ▼
RecommendationResponse {
  primary_recommendation, alternatives, summary, metadata
}
""")
P(doc, "Important detail: scoring is deterministic. The LLM only "
       "writes the human-readable summary. If the LLM is unavailable "
       "we still return ranked options — the user-facing failure mode "
       "is 'no nice paragraph', not 'no recommendation'.")

H2(doc, "12.4 Tracking Flow (Python API)")
CODE(doc, r"""
Frontend AdvisorPage (tracking tab)
  └─ POST /api/v1/advisor/tracking  { issue, context }
        │
        ▼
app/api/routes/advisor.py:69  tracking_advisor()
  - rag           = request.app.state.rag
  - tool_registry = request.app.state.tool_registry
  - llm_router    = request.app.state.llm_router
  - reasoning_client = llm_router.for_task(TASK_REASONING)   # → OpenAI
        │
        ▼
app/services/tracking_advisor_service.py  get_tracking_guidance()
  ├─ embedding_provider.embed([issue])
  ├─ vector_store.search(qvec, top_k=settings.rag_top_k)
  │     (retrieves delays/exceptions + claims/returns docs)
  ├─ if context has address dict:
  │     tool_registry.get("validate_address").execute(...)
  │     (e.g. checking the delivery address that 'failed delivery')
  ├─ build_advisor_prompt(issue, retrieved_chunks, tool_results)
  └─ reasoning_client.complete(messages)
        │
        ▼
TrackingAdvisorResponse {
  guidance,
  issue_summary,
  tools_used,
  sources [{source, chunk_index, score}],
  next_steps
}

Note: this flow shares almost everything with /advisor/shipping —
same retrieval, same reasoning task, same prompt builder. The only
difference is the service function and the response shape (guidance
+ next_steps instead of answer + reasoning_summary). That symmetry
is intentional and is the easiest interview answer to 'why isn't
this duplicated?' — both endpoints are thin wrappers around the
same RAG + tool + LLM kernel.
""")

PAGEBREAK(doc)


# ─────────────────────────────────────────────────────────────────────
# 13. FAILURE MODES & FALLBACK MATRIX
# ─────────────────────────────────────────────────────────────────────
H1(doc, "13. Failure Modes & Fallback Behaviour")
P(doc, "For each failure mode: what triggers it, what the system "
       "actually does, what the user sees, and what to verify.")

ftbl = doc.add_table(rows=1, cols=4)
ftbl.style = "Light Grid Accent 1"
hdr = ftbl.rows[0].cells
hdr[0].text = "Failure"
hdr[1].text = "Trigger"
hdr[2].text = "System behaviour"
hdr[3].text = "User-visible result"

failures = [
    ("Missing OpenAI key",
     "LLM_PROVIDER_REASONING=openai but OPENAI_API_KEY empty",
     "build_provider_client returns None → router falls through to LLM_PROVIDER_FALLBACK (echo). Warning logged: 'Provider openai requested but OPENAI_API_KEY is not set'.",
     "Advisor still answers but with EchoClient — returns the retrieved context verbatim with a note. Synthesis (Gemini) unaffected."),
    ("Missing Gemini key",
     "LLM_PROVIDER_SYNTHESIS=gemini but GEMINI_API_KEY empty",
     "Same path as above for the synthesis task only. Reasoning (OpenAI) unaffected.",
     "/rag/query and recommendation summary use Echo. /advisor/shipping still uses OpenAI."),
    ("Both LLM keys missing",
     "LLM_PROVIDER set but no keys, or LLM_PROVIDER unset entirely",
     "Both tasks degrade to LLM_PROVIDER_FALLBACK (echo). App boots normally — never crashes.",
     "All AI endpoints return retrieved-context responses. Frontend works; just no AI prose."),
    ("Missing carrier credentials",
     "SHIPPING_PROVIDER=ups with empty UPS_CLIENT_ID/SECRET",
     "_has_required_credentials() returns False → factory falls back to MockShippingProvider. Warning logged.",
     "Quote previews and address validation return mock data. Tools still work end-to-end."),
    ("Empty vector store",
     "Service started but /rag/ingest never called, or data/documents empty",
     "vector_store.search() returns []. RAG service passes empty context list to build_rag_prompt, which inserts 'No context was retrieved'. LLM is told to say it doesn't know.",
     "Advisor answers truthfully with 'no information available'. No hallucinated answer."),
    ("Python API down",
     "FastAPI process crashed, port unreachable",
     "Frontend fetch fails. Java API is unaffected — quoting and saved options still work.",
     "Advisor card shows error state. Quote form continues working normally."),
    ("RAG returns no matching chunks",
     "Query is unrelated to any document",
     "Top-K still returns the K closest vectors but with low scores. Prompt includes them; system prompt instructs grounding.",
     "Model typically responds 'the provided context does not cover this'. Sources panel shows the low-score chunks."),
    ("LLM provider call raises mid-request",
     "Network error, 5xx, rate limit",
     "OpenAIClient/GeminiClient catch and raise AppError(502). Global error handler returns clean JSON.",
     "Frontend sees a 502 with a message. Other endpoints unaffected."),
    ("Tool execution fails",
     "Provider returned an error or input failed validate_input",
     "AppError raised; route returns it. tools_used reflects what ran.",
     "Advisor returns the LLM answer derived from RAG context only, minus the failed tool's data."),
    ("Bad LLM_PROVIDER_FALLBACK name",
     "LLM_PROVIDER_FALLBACK=not-a-real-provider",
     "build_provider_client returns None; router catches and uses EchoClient. Tested in test_router_never_crashes_on_bad_fallback_name.",
     "Same as Echo fallback. App still boots."),
]
for f in failures:
    cells = ftbl.add_row().cells
    for i, v in enumerate(f):
        cells[i].text = v

doc.add_paragraph()
P(doc, "Design principle behind every row above: degrade in place, "
       "never crash on startup, always log the reason. The app boots "
       "with an entirely empty .env.")

PAGEBREAK(doc)


# ─────────────────────────────────────────────────────────────────────
# 14. SPRING BOOT — DEEPER DIVE
# ─────────────────────────────────────────────────────────────────────
H1(doc, "14. Spring Boot Backend — Deeper Dive")

H2(doc, "14.1 Request Lifecycle (one quote round-trip)")
CODE(doc, r"""
HTTP POST /api/v1/quotes
        │
        ▼
┌────────────────────────────────────────┐
│ Tomcat / Embedded Servlet Container    │
└────────────────────────────────────────┘
        │
        ▼
┌────────────────────────────────────────┐
│ Spring SecurityFilterChain             │
│  - JwtAuthenticationFilter             │
│  - extracts Authorization: Bearer ...  │
│  - validates against Supabase JWKS     │
│  - sets SecurityContextHolder          │
└────────────────────────────────────────┘
        │
        ▼
┌────────────────────────────────────────┐
│ DispatcherServlet → @RestController    │
│  QuoteController.getQuote(@Valid dto)  │
│   - Bean Validation runs on dto        │
│   - @PreAuthorize checks role/scope    │
└────────────────────────────────────────┘
        │
        ▼
┌────────────────────────────────────────┐
│ @Service QuoteService                  │
│  - business rules                      │
│  - @Transactional opens TX             │
│  - calls repository                    │
└────────────────────────────────────────┘
        │
        ▼
┌────────────────────────────────────────┐
│ Spring Data JPA Repository             │
│  - generated SQL via Hibernate         │
│  - persists audit row                  │
└────────────────────────────────────────┘
        │
        ▼
┌────────────────────────────────────────┐
│ TX commits → DTO mapper → JSON         │
│ HTTP 200 + QuoteResponse               │
└────────────────────────────────────────┘

Errors:
  Validation  → @ControllerAdvice → 400 + field errors
  Auth        → SecurityFilterChain → 401/403
  Not found   → custom exception → 404
  Anything else → fallback handler → 500 (logged)
""")

H2(doc, "14.2 DTO Validation")
P(doc, "Inputs are Java records / classes with Bean Validation "
       "annotations (@NotNull, @Positive, @Size, @Pattern). The "
       "controller method parameter is annotated @Valid, which makes "
       "Spring run the validators before the controller body executes. "
       "Validation failures throw MethodArgumentNotValidException, "
       "which a @ControllerAdvice maps to a 400 with a structured "
       "field-error list. The service layer never sees an invalid "
       "DTO.")

H2(doc, "14.3 Service Layer")
P(doc, "Services hold the business rules. They are stateless beans, "
       "constructor-injected with their dependencies (repositories, "
       "other services). Why a layer at all instead of putting logic "
       "in controllers? Three reasons:")
BULLET(doc, "Controllers stay HTTP-shaped (status codes, headers); services stay domain-shaped. Different testing strategies.")
BULLET(doc, "@Transactional belongs at the service boundary, not the controller, because one HTTP request might invoke several services.")
BULLET(doc, "Reuse: a scheduled job or a background worker can call the same service without going through HTTP.")

H2(doc, "14.4 Repository / Persistence")
P(doc, "Repositories extend JpaRepository<Entity, Id>. Spring Data "
       "generates implementations from method names "
       "(findByUserIdOrderByCreatedAtDesc) and gives you crudby and "
       "page-by-x for free. Hibernate is the JPA provider; the actual "
       "DB is Supabase Postgres. Entities are mapped with @Entity / "
       "@Id / @Column. The repository layer is the only place that "
       "talks SQL — everything above thinks in domain objects.")

H2(doc, "14.5 JWT / Auth Flow")
CODE(doc, r"""
Frontend
  - User signs in via Supabase Auth (handled in the SPA)
  - Receives a signed JWT (RS256)
  - Sends Authorization: Bearer <jwt> on every API call

Java API
  - SecurityFilterChain registers JwtAuthenticationFilter
  - Filter calls JwtDecoder.decode(token)
       (decoder fetches Supabase JWKS; verifies signature + exp + iss)
  - Builds an Authentication with sub/email/roles claims
  - Stores it in SecurityContextHolder
  - Controller methods can use @PreAuthorize and @AuthenticationPrincipal
  - Saved-options + bookings are scoped to the authenticated user id
""")

H2(doc, "14.6 Transaction Boundaries")
P(doc, "@Transactional is applied at the service method level. "
       "Default propagation is REQUIRED — if a transaction is already "
       "open it joins, otherwise it starts one. Read-only methods are "
       "annotated @Transactional(readOnly = true) so Hibernate skips "
       "dirty checking. Anything that mutates more than one row "
       "(saving a quote + writing an audit row) lives in one method "
       "so the whole thing is atomic. Crucially: never put "
       "@Transactional on the controller — controllers may call "
       "multiple services, and you want each to manage its own boundary.")

PAGEBREAK(doc)


# ─────────────────────────────────────────────────────────────────────
# 15. BACKEND DESIGN PATTERNS USED HERE
# ─────────────────────────────────────────────────────────────────────
H1(doc, "15. Backend Design Patterns Used Here")
P(doc, "Naming the patterns out loud is a senior-interview move. "
       "Each one below is paired with the exact file in this repo "
       "where it lives, so the answer is never abstract.")

H2(doc, "15.1 Factory")
P(doc, "create_llm_client(), create_llm_router(), create_shipping_provider(), "
       "create_embedding_provider(), create_vector_store(). Each one "
       "reads config and returns the right concrete implementation. "
       "Callers depend on the ABC, not the constructor of any "
       "specific class. Adding a new provider = one branch in one factory.")

H2(doc, "15.2 Adapter")
P(doc, "UPSProvider, FedExProvider, DHLProvider, USPSProvider all "
       "adapt their carrier's HTTP/JSON shape into the project's "
       "ShippingProvider interface (validate_address, "
       "get_quote_preview). The rest of the system never sees a "
       "carrier-specific field name. GeminiClient is also an adapter "
       "— it converts OpenAI-style {role, content} messages into "
       "Gemini's contents/parts format inside _messages_to_gemini_contents.")

H2(doc, "15.3 Strategy")
P(doc, "LLMRouter is strategy with a twist: instead of one strategy "
       "per request, it picks one strategy per task at startup and "
       "reuses it. The 'algorithm' you're swapping is which LLM "
       "client handles which logical task. Same shape — interface + "
       "interchangeable implementations selected by config.")

H2(doc, "15.4 Graceful Degradation")
P(doc, "Every factory has a fallback path. Every external call has "
       "a safe default. Examples already cited: bad creds → mock "
       "provider; missing key → echo; bad fallback name → echo; "
       "empty vector store → 'no context' branch in the prompt. The "
       "system can boot with an empty .env and serve requests.")

H2(doc, "15.5 Composition Root")
P(doc, "main.py's lifespan() function is the composition root. It "
       "is the only place that constructs concrete classes and wires "
       "them together. Every other file works against ABCs and "
       "receives its dependencies. This makes tests trivial: the "
       "test just builds its own app.state with fakes.")

H2(doc, "15.6 Abstraction Boundaries")
BULLET(doc, "Routes know about HTTP and Pydantic schemas, not LLMs.")
BULLET(doc, "Services know about LLMClient and ToolRegistry, not OpenAI/Gemini specifics.")
BULLET(doc, "Clients know about one provider each, not about routing or tasks.")
BULLET(doc, "The router knows about tasks and clients, not about HTTP or RAG.")
P(doc, "Each layer has exactly one reason to change. That is the "
       "single-responsibility principle in practice, and it is what "
       "makes this codebase explainable in 60 seconds.")

PAGEBREAK(doc)


# ─────────────────────────────────────────────────────────────────────
# 16. DESIGN TRADEOFFS
# ─────────────────────────────────────────────────────────────────────
H1(doc, "16. Design Tradeoffs")
P(doc, "For each decision: what we picked, the alternative we "
       "rejected, and the cost we accepted by picking it. Senior "
       "interviewers care more about the tradeoff than the choice.")

H2(doc, "16.1 Why the Frontend Orchestrates Two Backends")
BULLET(doc, "Picked: frontend calls Java and Python directly, merges results in the UI.")
BULLET(doc, "Alternative: a BFF / API gateway sits in front, fanning out to both.")
BULLET(doc, "Cost accepted: the frontend has to know two base URLs and two error shapes. We do not have a single auth/rate-limit chokepoint.")
BULLET(doc, "Reason: it removes one deployable, eliminates an internal RPC channel that could fail, and makes failure modes obvious to the user (one card errors, the rest of the page works).")

H2(doc, "16.2 Why Java and Python Are Split")
BULLET(doc, "Picked: Spring Boot for transactions; FastAPI for AI.")
BULLET(doc, "Alternative: one polyglot service (Java with a Python sidecar) or one all-Python service.")
BULLET(doc, "Cost accepted: two CI pipelines, two language ecosystems, two deploys.")
BULLET(doc, "Reason: each service can use its strongest ecosystem (Spring for transactional, Python for LLM/embedding libs) and a flaky AI dependency cannot bring down the booking flow.")

H2(doc, "16.3 Why Deterministic Scoring Before LLM Explanation")
BULLET(doc, "Picked: rank options in pure code, then ask the LLM to write a paragraph about the top one.")
BULLET(doc, "Alternative: ask the LLM to do the ranking too.")
BULLET(doc, "Cost accepted: the prose is constrained to 'explain this winner' instead of 'choose freely'.")
BULLET(doc, "Reason: ranking is the part the user can verify against price/ETA columns. Hallucinated rankings would erode trust immediately. LLMs are great writers and bad spreadsheets — give them the writing job.")

H2(doc, "16.4 Why an In-Memory Vector Store First")
BULLET(doc, "Picked: InMemoryVectorStore behind a VectorStore ABC.")
BULLET(doc, "Alternative: pgvector / Qdrant / FAISS from day one.")
BULLET(doc, "Cost accepted: data is lost on restart; ingestion is a route, not a job; performance ceiling around ~10k chunks.")
BULLET(doc, "Reason: the bottleneck of a RAG system is rarely the index at this scale. Putting an ABC in place means the swap is one factory edit when the corpus grows. We chose to demonstrate the architecture, not over-engineer the index.")

H2(doc, "16.5 Why Task-Based LLM Routing")
BULLET(doc, "Picked: a router that maps logical tasks (reasoning, synthesis, fallback) to clients, configured per env var.")
BULLET(doc, "Alternative: one global LLM client; or per-request dynamic selection with scoring.")
BULLET(doc, "Cost accepted: two SDK surfaces, two failure modes, slight tone variation between tasks.")
BULLET(doc, "Reason: cost (Gemini Flash for synthesis is ~10× cheaper), failure isolation, and operability — ops change a single env var to swap a task's provider without redeploying code.")

H2(doc, "16.6 Why Tools Instead of Pure LLM Reasoning")
BULLET(doc, "Picked: deterministic tools for ground-truth questions; LLM only for advice/synthesis.")
BULLET(doc, "Alternative: ask the LLM to compute everything from training-time knowledge.")
BULLET(doc, "Cost accepted: more code (Tool ABC, registry, schemas) and more handoffs.")
BULLET(doc, "Reason: LLMs hallucinate facts. ZIP validation and quote pricing are exactly the questions where 'sounds right' is not good enough. Tools make those answers boring and correct.")

PAGEBREAK(doc)


# ─────────────────────────────────────────────────────────────────────
# 17. DEBUGGING & VALIDATION
# ─────────────────────────────────────────────────────────────────────
H1(doc, "17. Debugging & Validation")
P(doc, "How to manually convince yourself each subsystem is working. "
       "Run these in order top to bottom — each step proves the next "
       "one's prerequisites.")

H2(doc, "17.1 Manual Test Sequence")
NUM(doc, "Boot the Python API with an empty .env. Expected log: 'LLM router initialized: {reasoning: echo, synthesis: echo, fallback: echo}'. Proves: factories degrade safely.")
NUM(doc, "GET /health → 200. Proves: process is up, routes mounted.")
NUM(doc, "GET /api/v1/orchestration/tools → list of 2 tool schemas. Proves: tool registry built.")
NUM(doc, "POST /api/v1/rag/ingest → response with chunks_ingested > 0. Proves: documents found, chunked, embedded, stored.")
NUM(doc, "POST /api/v1/rag/query with a known-content query → answer + sources array non-empty. Proves: retrieval + prompt + LLM (or Echo) path.")
NUM(doc, "POST /api/v1/advisor/shipping with a real query → answer + tools_used + sources. Proves: reasoning task wired, services compose retrieval + tools + LLM.")
NUM(doc, "POST /api/v1/advisor/recommendation with sample services → primary_recommendation + alternatives + summary. Proves: deterministic scoring runs even when LLM is Echo.")
NUM(doc, "Set OPENAI_API_KEY and LLM_PROVIDER_REASONING=openai, restart. Re-run /advisor/shipping. Log should now show provider=openai. Proves: live LLM path.")
NUM(doc, "Unset OPENAI_API_KEY but keep LLM_PROVIDER_REASONING=openai. Restart. Expected warning + reasoning falls to fallback. Proves: per-task degradation.")

H2(doc, "17.2 Logs to Inspect")
BULLET(doc, "Startup: 'Starting shipsmart-api-python ...', 'LLM router initialized: {...}', 'Tool registry initialized: 2 tools, provider=mock'. If any of these are missing, that subsystem failed to wire.")
BULLET(doc, "Per-task resolution: 'LLM router: task=reasoning → provider=openai'. Proves which provider each task picked.")
BULLET(doc, "Warnings: 'Provider openai requested but OPENAI_API_KEY is not set' or 'Failed to create provider ups: ... — falling back to mock'. These are the smoking guns for misconfig.")
BULLET(doc, "Per-request: RequestLoggingMiddleware logs the method, path, status, and duration. Useful for spotting slow LLM calls.")
BULLET(doc, "Errors: AppError-derived exceptions are formatted by the global handler; 502s typically mean an LLM provider returned an error.")

H2(doc, "17.3 What 'Working' Looks Like Per Subsystem")
ttbl = doc.add_table(rows=1, cols=3)
ttbl.style = "Light Grid Accent 1"
hdr = ttbl.rows[0].cells
hdr[0].text = "Subsystem"
hdr[1].text = "How to exercise it"
hdr[2].text = "Output that proves it works"
proofs = [
    ("Config", "boot with .env", "No exceptions; defaults applied for unset vars."),
    ("Lifespan / wiring", "boot logs", "All five 'initialized' lines present."),
    ("Tool registry", "GET /api/v1/orchestration/tools", "JSON list with validate_address + get_quote_preview schemas."),
    ("Shipping provider", "GET /api/v1/orchestration/tools or run a tool", "Provider name in startup log; tool execution returns sane data."),
    ("Embeddings", "POST /rag/ingest", "chunks_ingested > 0; no zip(strict=True) error."),
    ("Vector store", "POST /rag/query", "sources array non-empty for an in-corpus query."),
    ("LLM router", "startup logs", "'LLM router initialized: {reasoning, synthesis, fallback}' line."),
    ("Reasoning task", "/advisor/shipping with key set", "Coherent answer; not the Echo template."),
    ("Synthesis task", "/rag/query with key set", "Coherent answer; not the Echo template."),
    ("Per-task fallback", "unset one key, restart, hit that endpoint", "Echo-style answer for that task only; the other still uses its real LLM."),
    ("Frontend integration", "AdvisorPage in browser", "Card renders sources and answer; loading + error states correct."),
]
for r in proofs:
    cells = ttbl.add_row().cells
    for i, v in enumerate(r):
        cells[i].text = v

PAGEBREAK(doc)


# ─────────────────────────────────────────────────────────────────────
# 18. INTERVIEW PRACTICE PACK
# ─────────────────────────────────────────────────────────────────────
H1(doc, "18. Interview Practice Pack")
P(doc, "Each block: question — strong sample answer in your own "
       "voice. Read them out loud at least twice. The point is not "
       "to memorise the wording, it is to internalise the structure: "
       "decision → tradeoff → evidence in code.")

H2(doc, "18.1 Senior Backend Questions")

H3(doc, "Q. Walk me through how a single HTTP request becomes a database row.")
P(doc, "On the Java side: the request hits Tomcat, then the Spring "
       "SecurityFilterChain validates the Supabase JWT and populates "
       "SecurityContextHolder. DispatcherServlet routes to the "
       "@RestController method, which receives a @Valid DTO — Bean "
       "Validation runs before the body executes. The controller "
       "calls a @Service method, which is where @Transactional opens "
       "the transaction. The service uses a Spring Data JPA "
       "repository to persist via Hibernate, the transaction commits, "
       "the response DTO is mapped, and Spring serialises it to JSON. "
       "Failures route through @ControllerAdvice into structured "
       "error responses.")

H3(doc, "Q. How do you handle errors and validation consistently?")
P(doc, "Validation lives at the boundary using Bean Validation "
       "annotations on DTOs, triggered by @Valid. A "
       "@ControllerAdvice maps validation exceptions to 400s with "
       "field-level details. Domain errors throw typed exceptions "
       "that the same advice maps to 4xx codes. Anything unexpected "
       "falls through to a 500 handler that logs the stack trace and "
       "returns a sanitised body. On the Python side the equivalent "
       "is AppError + register_error_handlers in app/core/errors.py.")

H3(doc, "Q. Why do you put @Transactional on the service and not the controller?")
P(doc, "Because the controller is HTTP-shaped and might call several "
       "services. Each service should own its own transactional "
       "boundary. Putting @Transactional on the controller means a "
       "transaction stays open across multiple service calls, which "
       "blurs ownership and makes rollback semantics fuzzy. It also "
       "couples HTTP concerns to persistence concerns.")

H3(doc, "Q. How would you scale this backend?")
P(doc, "Java side: it is already stateless, so horizontal scaling "
       "behind a load balancer is the first move. Push read-heavy "
       "queries through a read replica. Add Redis for hot data. "
       "Python side: same — stateless apart from app.state caches "
       "that are easy to externalise. The only shared state today is "
       "the in-memory vector store, and that is behind an ABC, so "
       "swapping to pgvector is a one-factory edit. The harder "
       "scaling problem is LLM cost — task-based routing already "
       "addresses that by sending the cheaper task to a cheaper "
       "provider.")

H2(doc, "18.2 RAG Questions")

H3(doc, "Q. Walk me through your RAG pipeline.")
P(doc, "On startup the service walks data/documents recursively, "
       "chunks each .md/.txt into ~500-character overlapping windows, "
       "embeds them — text-embedding-3-small in production, a "
       "deterministic hash embedding in dev — and stores them in an "
       "in-memory vector store behind an ABC. At query time we embed "
       "the query, take cosine top-3, and call build_rag_prompt() "
       "which puts the chunks in a system-grounded prompt that tells "
       "the LLM to answer only from the provided context. The "
       "endpoint returns the answer plus the sources (file path, "
       "chunk index, score) so the UI can cite them.")

H3(doc, "Q. How do you avoid hallucination?")
P(doc, "Five layers. First, retrieval: we put real document text in "
       "the prompt so the model has something to ground on. Second, "
       "system prompt: it explicitly says answer only from the "
       "provided context. Third, tools: ground-truth questions "
       "(addresses, quotes) go through deterministic tools, never "
       "the LLM. Fourth, citations: every answer comes with sources, "
       "so the user can verify. Fifth, fallback: if the LLM "
       "is unavailable we return retrieved context verbatim via "
       "EchoClient instead of an empty hallucination.")

H3(doc, "Q. Why chunking with overlap? Why not embed full documents?")
P(doc, "Two reasons. One, the LLM has a context limit; you can't "
       "stuff a whole policy doc into every prompt. Two, embeddings "
       "average semantics across the input — a long document "
       "produces a vector that represents nothing in particular. "
       "Chunking gives each section its own vector, so the retrieval "
       "step can rank passages independently. The overlap exists so "
       "a sentence that straddles a chunk boundary still appears "
       "intact in at least one chunk.")

H3(doc, "Q. What would change at 10 million documents?")
P(doc, "Index, ingestion, and embedding costs. The InMemoryVectorStore "
       "would be replaced with pgvector or Qdrant — same ABC, swap "
       "the factory. Ingestion would move out of the request path "
       "into a background job. We'd batch and cache embedding calls. "
       "We'd add metadata filters (carrier, region) so retrieval "
       "narrows the candidate set before ranking. None of this "
       "requires touching the services or routes — that's the whole "
       "point of the abstraction.")

H2(doc, "18.3 MCP / Tooling Questions")

H3(doc, "Q. Why do you have a tool layer at all? Why not let the LLM answer?")
P(doc, "Because some questions have a right answer and the LLM "
       "doesn't know it. 'Is this a valid US ZIP code?' and 'how "
       "much will UPS charge for this package?' should be computed "
       "by code that talks to a real API, not predicted from "
       "training data. The tool layer is where ground truth lives. "
       "The LLM stays in the lane it is good at — explaining the "
       "answer in natural language.")

H3(doc, "Q. How is your tool layer like MCP?")
P(doc, "Same shape. Each tool exposes name, description, "
       "parameters with types, and an execute method. The registry "
       "lists schemas via /api/v1/orchestration/tools — that is the "
       "discovery surface. Inputs are validated before execution. "
       "The handoff between text reasoning and structured execution "
       "is what MCP standardises, and we model it the same way "
       "internally.")

H2(doc, "18.4 Weak vs Strong Answer Examples")
P(doc, "Same question, two answers. Read both and notice the "
       "difference: the strong one names a file, a tradeoff, and an "
       "alternative. The weak one is a generic talking point that "
       "could come from anyone's blog.")

H3(doc, "Q. How does your RAG pipeline avoid hallucination?")
P(doc, "Weak: 'We use grounding and citations. The model only "
       "answers from the retrieved context, which prevents "
       "hallucination.'")
P(doc, "Strong: 'Five layers, in order. One, retrieval — "
       "rag_service.py embeds the query and pulls top-3 chunks "
       "from the in-memory store. Two, the system prompt in "
       "build_rag_prompt() literally says answer ONLY from the "
       "provided context. Three, ground-truth questions like "
       "address validity bypass the LLM entirely and go through "
       "ValidateAddressTool. Four, sources come back in the "
       "response with file path and similarity score so the UI "
       "can cite them — the user can verify any answer. Five, if "
       "the LLM is unavailable EchoClient returns the retrieved "
       "context verbatim instead of a confident-sounding "
       "fabrication. The tradeoff we accept is that the model "
       "sometimes says ''the provided context does not cover "
       "this'', which is the right failure mode.'")

H3(doc, "Q. Why split Java and Python?")
P(doc, "Weak: 'Java is good for backend and Python is good for "
       "AI, so we used the right tool for each job.'")
P(doc, "Strong: 'Ownership and failure isolation, not language "
       "preference. Spring Boot is the system of record for quotes, "
       "saved options, and bookings — JPA, @Transactional, "
       "validated DTOs, JWT auth, the boring mature path. FastAPI "
       "owns the AI layer because the LLM/embedding ecosystem is "
       "Python-first and async I/O for OpenAI/Gemini calls is "
       "natural in Starlette. Crucially, the two backends never "
       "call each other — the frontend orchestrates. That means a "
       "flaky LLM provider cannot block a quote being saved, and "
       "the AI service can deploy on its own cycle. The cost we "
       "accepted is two CI pipelines and two deploy artefacts. The "
       "alternative I rejected was a BFF in front of both — that "
       "adds a single point of failure and gives no benefit at our "
       "scale.'")

H3(doc, "Q. Why do you have a tool layer instead of letting the LLM answer everything?")
P(doc, "Weak: 'Tools make the LLM more reliable and prevent "
       "hallucination on factual questions.'")
P(doc, "Strong: 'Some questions have a right answer the LLM "
       "doesn't know. ''Is 90211 a valid Beverly Hills ZIP?'' and "
       "''what does UPS Ground cost from Boston to Seattle for a "
       "5lb package?'' should be computed by code that hits a real "
       "API, not predicted from training data. So we built a Tool "
       "ABC with name, description, parameters and async execute, "
       "registered in a ToolRegistry, discoverable via "
       "/api/v1/orchestration/tools — same shape as MCP. "
       "ValidateAddressTool and GetQuotePreviewTool both delegate "
       "to the configured ShippingProvider, so the same code paths "
       "work against the mock provider in dev and against UPS in "
       "production. Tool outputs are appended to the LLM prompt as "
       "evidence so the model''s answer is grounded. The LLM stays "
       "in the lane it''s good at: writing prose that explains the "
       "tool result.'")

H3(doc, "Q. How would you scale this system?")
P(doc, "Weak: 'I'd add caching, horizontal scaling, and a real "
       "vector database.'")
P(doc, "Strong: 'Different layers scale differently. The Java "
       "side is already stateless behind a load balancer; the "
       "first move is read replicas for the quote-history queries. "
       "The Python side is stateless apart from app.state, which "
       "is just prebuilt clients — easy to replicate. The actual "
       "bottlenecks I''d hit first are LLM cost and vector store "
       "size. For LLM cost, task-based routing already sends the "
       "cheaper task to Gemini Flash; the next move is per-task "
       "model overrides so I can drop the synthesis temperature "
       "and max_tokens independently. For vector store, "
       "InMemoryVectorStore is behind an ABC so swapping to "
       "pgvector or Qdrant is one factory edit — but I''d only do "
       "it once the corpus is past ~10k chunks, because brute-"
       "force cosine is fine below that and the index is rarely "
       "the real bottleneck. The bigger win is moving ingestion "
       "out of the request path into a background job and caching "
       "embeddings of common queries.'")

H3(doc, "Q. How do you keep tool outputs and LLM reasoning consistent?")
P(doc, "Tool outputs are appended to the LLM prompt as evidence "
       "before the model writes its answer. The system prompt "
       "instructs the model to use the tool data as ground truth. "
       "The response back to the client includes both the LLM text "
       "and a tools_used list, so the UI can show what was computed "
       "vs. what was reasoned. If the user sees a price in the "
       "answer, they also see which tool produced it.")

PAGEBREAK(doc)


# ─────────────────────────────────────────────────────────────────────
# 19. MUST KNOW vs NICE TO KNOW
# ─────────────────────────────────────────────────────────────────────
H1(doc, "19. Prioritisation: Must / Should / Nice to Know")
P(doc, "Three explicit tiers. MUST KNOW is the floor — if any of "
       "these are shaky, stop learning new material and rehearse "
       "until they aren't. SHOULD KNOW is what turns a competent "
       "answer into a senior one. NICE TO KNOW only matters if the "
       "interviewer drills in deeper than expected.")

mtbl = doc.add_table(rows=1, cols=3)
mtbl.style = "Light Grid Accent 1"
hdr = mtbl.rows[0].cells
hdr[0].text = "MUST KNOW (floor)"
hdr[1].text = "SHOULD KNOW (senior signal)"
hdr[2].text = "NICE TO KNOW (depth)"

mrows = [
    ("3-service architecture; Java for transactions, Python for AI; frontend orchestrates",
     "Why the two backends never call each other (failure isolation, independent deploy)",
     "Comparison with BFF / API gateway alternatives"),
    ("RAG = retrieve, then ground the LLM in the retrieved text",
     "Chunking with overlap; cosine top-K; sources returned for citation",
     "Hybrid retrieval (BM25 + dense), reranking, eval metrics like recall@k"),
    ("Tools exist to handle ground-truth questions instead of letting the LLM guess",
     "Tool ABC + registry + schemas; how it mirrors MCP",
     "Function calling vs explicit tool dispatch; tool authorisation"),
    ("Task-based LLM router: reasoning → OpenAI, synthesis → Gemini, fallback → Echo",
     "Per-task resolution order; per-task independent fallback; cost rationale",
     "Per-task tuning (timeout/temperature), provider-side retries, A/B routing"),
    ("Composition root in main.py lifespan(); everything else uses ABCs",
     "Why building once at startup matters (cost, testability)",
     "Dependency injection containers vs hand-wired roots"),
    ("Spring Boot request lifecycle: filter chain → controller → service → repo",
     "@Transactional on the service layer, not the controller",
     "Propagation modes, isolation levels, optimistic locking"),
    ("DTO validation with @Valid + Bean Validation; @ControllerAdvice for errors",
     "Why validation must live at the boundary",
     "Custom validators, group validation, cross-field constraints"),
    ("JWT auth from Supabase, validated in Spring SecurityFilterChain",
     "JWKS rotation, claims-based authorization with @PreAuthorize",
     "Token refresh strategies, clock skew, audience claims"),
    ("Graceful degradation: every factory has a safe default",
     "Specific fallbacks: bad creds → mock; missing key → echo",
     "Circuit breakers, bulkheads, exponential backoff per provider"),
    ("Deterministic scoring before LLM explanation in recommendations",
     "Why ranking is in code and only the prose is in the LLM",
     "Learning-to-rank, ranking models, score calibration"),
]
for r in mrows:
    cells = mtbl.add_row().cells
    for i, v in enumerate(r):
        cells[i].text = v

doc.add_paragraph()
P(doc, "Strategy for the day before: re-read column 1 once, then "
       "do every Out-Loud item in §10. If anything in column 1 is "
       "shaky, stop studying new material and rehearse it.")


# ── Save ──────────────────────────────────────────────────────────
out = "docs/assets/ShipSmart-Study-Guide-v3.docx"
doc.save(out)
print(f"Wrote {out}")
